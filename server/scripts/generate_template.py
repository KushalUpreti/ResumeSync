from docxtpl import DocxTemplate
from docx import Document
import os

def create_base_template():
    # Create a fresh document
    doc = Document()
    
    # Add Placeholders
    doc.add_heading('{{ name }}', 0)
    doc.add_paragraph('{{ email }} | {{ phone }} | {{ location }}')
    
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('{{ summary }}')
    
    doc.add_heading('Professional Experience', level=1)
    # Loop for experience
    doc.add_paragraph('{% for item in experience %}')
    doc.add_heading('{{ item.role }} @ {{ item.company }}', level=2)
    doc.add_paragraph('{% if item.start_date or item.end_date %}{{ item.start_date }}{% if item.start_date and item.end_date %} - {% endif %}{{ item.end_date }}{% endif %}')
    doc.add_paragraph('{% for bullet in item.bullets %}')
    doc.add_paragraph('• {{ bullet }}', style='List Bullet')
    doc.add_paragraph('{% endfor %}')
    doc.add_paragraph('{% endfor %}')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('{{ skills|join(", ") }}')
    
    # Save it
    output_path = 'app/static/templates/modern.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Template created at {output_path}")

if __name__ == "__main__":
    create_base_template()
