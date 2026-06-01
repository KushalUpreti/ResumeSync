from __future__ import annotations

import io
from docxtpl import DocxTemplate
from app.models.resume import ResumeDocument
from app.services.interfaces import DocumentRenderer

class DocxtplDocumentRenderer(DocumentRenderer):
    def __init__(self, template_dir: str = "app/static/templates") -> None:
        self.template_dir = template_dir

    def render(self, document: ResumeDocument, *, template_id: str) -> bytes:
        # For now, we'll assume the template exists locally in the container/server
        # In a more advanced setup, we could fetch templates from S3
        template_path = f"{self.template_dir}/{template_id}.docx"
        
        try:
            doc = DocxTemplate(template_path)
        except Exception:
            # Fallback if template doesn't exist - create a very basic doc
            from docx import Document
            fallback_doc = Document()
            fallback_doc.add_heading(document.full_name, 0)
            fallback_doc.add_paragraph(f"{document.email} | {document.phone} | {' | '.join(document.links)}")
            fallback_doc.add_paragraph(document.summary)
            
            fallback_doc.add_heading("Experience", level=1)
            for exp in document.experience:
                fallback_doc.add_heading(f"{exp.role} at {exp.company}", level=2)
                date_parts = [part for part in [exp.start_date, exp.end_date] if part]
                if date_parts:
                    fallback_doc.add_paragraph(" | ".join(date_parts))
                for bullet in exp.bullets:
                    fallback_doc.add_paragraph(bullet, style='List Bullet')
            
            fallback_doc.add_heading("Skills", level=1)
            for skill_cat in document.skills:
                fallback_doc.add_heading(skill_cat.category, level=2)
                fallback_doc.add_paragraph(", ".join(skill_cat.items))

            if document.education:
                fallback_doc.add_heading("Education", level=1)
                for edu in document.education:
                    title_parts = [edu.degree]
                    if edu.field_of_study:
                        title_parts.append(f"in {edu.field_of_study}")
                    fallback_doc.add_heading(f"{' '.join(title_parts)} — {edu.institution}", level=2)
                    date_parts = [part for part in [edu.start_date, edu.end_date] if part]
                    if date_parts:
                        fallback_doc.add_paragraph(" – ".join(date_parts))
                    if edu.gpa:
                        fallback_doc.add_paragraph(f"GPA: {edu.gpa}")
                    if edu.description:
                        fallback_doc.add_paragraph(edu.description)

            if document.projects:
                fallback_doc.add_heading("Projects", level=1)
                for proj in document.projects:
                    title_parts = [proj.name]
                    if proj.role:
                        title_parts.append(proj.role)
                    fallback_doc.add_heading(" — ".join(title_parts), level=2)
                    date_parts = [part for part in [proj.start_date, proj.end_date] if part]
                    if date_parts:
                        fallback_doc.add_paragraph(" – ".join(date_parts))
                    if proj.url:
                        fallback_doc.add_paragraph(f"URL: {proj.url}")
                    if proj.technologies:
                        fallback_doc.add_paragraph(f"Technologies: {', '.join(proj.technologies)}")
                    if proj.description:
                        fallback_doc.add_paragraph(proj.description)
                    for bullet in proj.bullets:
                        fallback_doc.add_paragraph(bullet, style='List Bullet')

            if document.certifications:
                fallback_doc.add_heading("Certifications", level=1)
                for cert in document.certifications:
                    parts = [cert.name]
                    if cert.issuer:
                        parts.append(cert.issuer)
                    if cert.date_obtained:
                        parts.append(cert.date_obtained)
                    fallback_doc.add_paragraph(" | ".join(parts))

            target = io.BytesIO()
            fallback_doc.save(target)
            return target.getvalue()

        # Context for docxtpl
        context = {
            "full_name": document.full_name,
            "email": document.email,
            "phone": document.phone,
            "links": document.links,
            "links_csv": " | ".join(document.links),
            "summary": document.summary,
            "experience": [
                {
                    "company": exp.company,
                    "role": exp.role,
                    "start_date": exp.start_date or "",
                    "end_date": exp.end_date or "",
                    "bullets": exp.bullets
                }
                for exp in document.experience
            ],
            "skills": [
                {
                    "category": cat.category,
                    "items": cat.items,
                    "skills_csv": ", ".join(cat.items),
                }
                for cat in document.skills
            ],
            "education": [
                {
                    "institution": edu.institution,
                    "degree": edu.degree,
                    "field_of_study": edu.field_of_study or "",
                    "start_date": edu.start_date or "",
                    "end_date": edu.end_date or "",
                    "gpa": edu.gpa or "",
                    "description": edu.description or "",
                }
                for edu in document.education
            ],
            "has_education": len(document.education) > 0,
            "projects": [
                {
                    "name": proj.name,
                    "description": proj.description or "",
                    "role": proj.role or "",
                    "technologies": proj.technologies,
                    "url": proj.url or "",
                    "start_date": proj.start_date or "",
                    "end_date": proj.end_date or "",
                    "bullets": proj.bullets,
                }
                for proj in document.projects
            ],
            "has_projects": len(document.projects) > 0,
            "certifications": [
                {
                    "name": cert.name,
                    "issuer": cert.issuer or "",
                    "date_obtained": cert.date_obtained or "",
                    "url": cert.url or "",
                }
                for cert in document.certifications
            ],
            "has_certifications": len(document.certifications) > 0,
        }

        doc.render(context)
        
        target = io.BytesIO()
        doc.save(target)
        return target.getvalue()
